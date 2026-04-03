from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

OXFORD = RGBColor(0x05, 0x05, 0x36)
PURPLE = RGBColor(0x64, 0x07, 0xFA)
TURQ   = RGBColor(0x34, 0xE1, 0xD2)
BLUE   = RGBColor(0x00, 0xAE, 0xFF)
GREY   = RGBColor(0x5A, 0x5F, 0x7A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)
ORANGE = RGBColor(0xC0, 0x55, 0x10)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = OXFORD
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '6407FA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = PURPLE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = OXFORD
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run('  ' + text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = ORANGE
    run.italic = True
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '050536')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = 'F5F7FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(chr(0x2500) * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ============================================================
# COVER
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
run = p.add_run('ATLAS')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('THE ZING OPERATING PLATFORM')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = OXFORD

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Full Scope  |  Tech Stack  |  Build Plan')
run.font.size = Pt(12)
run.font.color.rgb = GREY
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by Max, Chief of Staff  |  March 2026  |  CONFIDENTIAL')
run.font.size = Pt(9)
run.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
h1('Executive Summary')

body(
    'Atlas is the central nervous system of the entire ZING operation. '
    'It replaces HubSpot and unifies every department into a single cloud-hosted system '
    'built for 100,000 customers and 100+ employees.'
)
body(
    'Beyond operations, Atlas is engineered as a data asset. Every customer interaction, '
    'website visit, financial event, and churn signal is captured and structured. '
    'The aggregate data across tens of thousands of small businesses becomes independently '
    'valuable and monetisable over time.'
)
body(
    'An AI intelligence layer runs continuously across all platform data, surfacing '
    'recommendations on sales performance, onboarding health, financial risk, customer '
    'engagement, and marketing effectiveness. Atlas turns operational data into a '
    'genuine business advisor.'
)

table(
    ['Metric', 'Value'],
    [
        ['Platform Name', 'Atlas'],
        ['Modules', '14'],
        ['Screens / Views', '~40'],
        ['Database Tables', '15+'],
        ['API Endpoints', '~60'],
        ['Employee Roles', '8'],
        ['External Integrations', '7'],
        ['Build Timeline', '6 weeks (4 weeks core + 2 weeks AI/data)'],
        ['Estimated Monthly Cost', '$25 to $160 (scales with usage)'],
    ],
    [2.2, 3.8]
)

# ============================================================
# PART 1: FULL SCOPE
# ============================================================
doc.add_page_break()
h1('Part 1: Full Scope')

body(
    'Atlas is structured across three strategic layers. All three are built together '
    'from day one. The data infrastructure and AI layer are not afterthoughts.',
    italic=True
)

table(
    ['Layer', 'Purpose'],
    [
        ['Operations', 'Run the business day to day: pipeline, onboarding, AR, support, campaigns, workflows'],
        ['Intelligence', 'AI continuously analyses all data and surfaces actionable recommendations'],
        ['Data', 'Capture and structure small business data that becomes a sellable asset'],
    ],
    [1.8, 4.2]
)

# MODULE 1
h2('Module 1: Dashboard')
body('The executive pulse. Everything Amy needs to see business health at a glance without clicking elsewhere.')

h3('Revenue Metrics')
for b in [
    'Period Revenue, Today\'s Revenue, NRR (Net Revenue Retention)',
    'Revenue by Deal Type: New vs Upgrade vs Add-on with % split and progress bars',
    'MRR per business day',
    'Monthly target vs actual with gap analysis and close-rate projection',
]:
    bullet(b)

h3('Visuals')
for b in [
    'Daily revenue bar chart. Click any bar to drill into that day\'s deals',
    'Interactive calendar. Click any date to see exactly what closed',
    'Sales rep leaderboard with individual breakdowns by product category',
]:
    bullet(b)

h3('AI Panel')
for b in [
    'Top 3 business insights surfaced automatically today',
    'Flagged anomalies (e.g. "Revenue down 18% vs same period last month")',
    'Recommended actions with one-click drill-down to supporting data',
]:
    bullet(b)

h3('Operational Summaries')
for b in [
    'AR health: current, past due, unpaid counts and dollar amounts',
    'Onboarding health: published this week, in progress, overdue items count',
]:
    bullet(b)

# MODULE 2
h2('Module 2: CRM / Contacts')
body('The master record of every person ZING has ever touched. Built for 100,000+ customers.')

h3('Contact List')
for b in [
    'Name, email, company, campaign source, lead source, status, lifetime value',
    'Search and filter (Active, Lead, Inactive)',
    'Sortable columns, bulk actions',
]:
    bullet(b)

h3('Contact Detail: Four Tabs')
for b in [
    'Customer Info: all fields, linked deals, linked tickets, business enrichment data (industry, revenue range, employee count, years in business, marketing budget)',
    'Pre-Sale Communications: chronological timeline of every touchpoint before close including campaign entry, outreach, calls, pipeline movements, and appointments',
    'Post-Sale Communications: timeline after close including onboarding completions, tickets, billing events, AR actions, and website launch',
    'Cancellation: cancellation request, retention attempts, win-back history, churn risk score and contributing factors',
]:
    bullet(b)

for b in [
    'Lead Sources tracked: Email, SMS, Paid Ads, Referral, Organic',
    'Every contact linked to the campaign that sourced them',
]:
    bullet(b)

# MODULE 3
h2('Module 3: Sales Pipeline')
body('The sales team\'s daily workspace. Each rep lives here.')

h3('Rep Tabs and Metrics')
for b in [
    'Individual rep tabs plus an All view for managers and Amy',
    'Date range filter for Won deals (default: This Month)',
    'Stat cards: Avg MRR per deal, MRR per business day, Won count, Active appointments',
    'Team leaderboard with bar rankings visible on the All tab',
    'Product breakdown: deals per product, total value',
]:
    bullet(b)

h3('Kanban Board: 9 Stages')
body('Call Now  >  Call No Answer  >  Hot 72  >  Active  >  Appointment  >  Appt No Show  >  Marketing Appt  >  Promotional  >  Won')
for b in [
    'Drag and drop between stages',
    'Deal cards show: product, value, priority, designer, rep (on All view)',
]:
    bullet(b)

h3('Won Deal Modal')
for b in [
    'Deal Type: New Customer, Upgrade, or Add-on',
    'Product and MRR amount',
    'Delivery date and assigned designer',
    'Launch Fee: single payment or split across multiple dates with allocation tracker and validation',
]:
    bullet(b)

h3('AI Coaching Nudges')
for b in [
    'Per-rep insights visible to that rep contextually inside the pipeline view',
    'Example: "Your hot-72 to appointment conversion is 12% below your 90-day average"',
    'Managers see team-wide pattern flags',
]:
    bullet(b)

# MODULE 4
h2('Module 4: Onboarding (Customer View)')
body('Tracks every customer\'s journey from sale to fully live. Built for onboarding managers.')

for b in [
    'Filters: assigned team member, status, due date urgency (overdue, due soon, upcoming)',
    'Stat cards: Total, In Progress, Not Started, Published, Average Progress %',
]:
    bullet(b)

h3('Customer Cards (expandable)')
for b in [
    'Circular progress indicator showing completion %',
    'Business details, contact info, website URLs',
    'Offshore designer and US designer assignment',
    'All 12 onboarding items with owner, status, due date',
    'AI at-risk badge for customers predicted to churn based on onboarding velocity',
]:
    bullet(b)

h3('The 12 Onboarding Items')
body('Website Design, AI Chat, Landing Pages, Blogs, Online Bookings, Memberships, Social Media, SMS Marketing, Email Marketing, GBP Optimization, Google Business Reviews, Local Directories')

h3('Website Design: 8-Stage Workflow')
table(
    ['Stage', 'Team'],
    [
        ['Not Started', 'None'],
        ['Website Started', 'Offshore Designer'],
        ['First Draft Ready', 'Offshore Designer'],
        ['Website Sent', 'US Designer'],
        ['Edits Mode', 'US Designer'],
        ['Ready for QA', 'US Designer'],
        ['Ready for Publishing', 'Publisher'],
        ['Published', 'Publisher'],
    ],
    [3.0, 3.0]
)

body('Each sub-stage has its own owner. Advancing a stage auto-suggests the correct team. Due dates are auto-calculated from sale date per item type. Overdue and due-soon items surface automatically.')

# MODULE 5
h2('Module 5: Onboarding (Task View)')
body('Same data, different lens. Instead of one customer\'s full journey, this shows all customers for one specific task. This is what designers and publishers use daily.')

for b in [
    '12 task tabs, one per onboarding item',
    'Per tab: every customer with that task, their status, owner, and due date',
    'Team workload bar showing active vs complete per person',
    'Stat cards: Complete, In Progress, Not Started, Overdue, Unassigned',
    'Inline status and owner updates without leaving the page',
]:
    bullet(b)

# MODULE 6
h2('Module 6: Support Tickets')
for b in [
    'Ticket list: ID, subject, contact, priority, category, date, status',
    'Side-panel detail view',
    'Status workflow: Open > In Progress > Resolved > Reopen',
    'Categories: Bug, Question, Feature Request, Billing',
    'AI pattern detection: when the same issue type appears 3+ times in 30 days, flagged as systemic rather than isolated (surfaced on Dashboard and Insights page)',
]:
    bullet(b)

# MODULE 7
h2('Module 7: Products')
body('Product catalogue and commission engine.')
for b in [
    'Product list: ID, description, price, category, commission structure, created date',
    'Categories: Subscription Monthly, Subscription Annual, One-Time',
    'Three commission models with live calculator: MRR Multiplier, % of One-Time Fee, % of Annual Contract Value',
    'Add, Edit, and Delete products',
]:
    bullet(b)

h3('Current Products')
table(
    ['Product', 'Price', 'Commission'],
    [
        ['DISCOVER', '$59/mo', '1x MRR'],
        ['BOOST', '$149/mo', '1x MRR'],
        ['DOMINATE', '$249/mo', '2x MRR'],
    ],
    [2.5, 1.5, 2.0]
)

# MODULE 8
h2('Module 8: Accounts Receivable')
body('Collections tracking and payment management. Replaces the Chase agent.')
for b in [
    'AR accounts with status: Current, Past Due, Unpaid, Paid',
    'Per account: business, customer, product, MRR, days past due, amount owed',
    'Full action timeline: Text Sent, Email Sent, Call Made, Stripe Retry, Escalated, Payment Received',
    'Manual timeline entry for logging calls and decisions',
    'AI churn risk score per account: Low, Medium, High, or Critical. Scored on payment history, onboarding completion, support tickets, and industry trends',
    'Stripe webhooks auto-update AR status when payment received',
]:
    bullet(b)

# MODULE 9
h2('Module 9: Campaigns')
body(
    'Marketing outreach management. Two views in one module: a campaign list for managing '
    'outreach programmes, and a pipeline funnel builder for automated stage-based SMS.'
)

h3('Campaigns List')
for b in [
    'Create and manage campaigns by type: SMS Blast, Email Blast, Cold Call, LinkedIn, Direct Mail, etc.',
    'Each campaign has its own contact list, CSV upload, and stat breakdown',
    'Stats per campaign: total contacts, CRM imports, active vs leads',
    'Date range filtering with presets (This Month, Last Month, Last 30 Days, YTD, All Time)',
    'Campaign detail view: click into any campaign to see all associated contacts',
    'Contact upload: CSV import tagged to campaign, validates name, email, phone, company columns',
]:
    bullet(b)

h3('Pipeline Funnels')
body(
    'Automated SMS messages that fire when a deal enters a specific pipeline stage. '
    'One funnel per stage, fully configurable.'
)
for b in [
    'One SMS automation per pipeline stage (Call No Answer, Hot 72, Appt No Show, Promotional)',
    'Toggle each funnel on or off independently',
    'Editable message body with merge fields: {{name}}, {{business_name}}, {{booking_link}}, {{promo_link}}',
    'Configurable send time and delay (in hours) before the message fires',
    'Default messages pre-loaded for each stage, editable at any time',
]:
    bullet(b)

h3('Pre-loaded Funnel Messages')
table(
    ['Stage', 'Default Timing', 'Message Purpose'],
    [
        ['Call No Answer', '9:00am, 24h delay', 'Re-engage after missed call, offer to schedule'],
        ['Hot 72', '10:00am, 24h delay', 'Urgency nudge, hold a design slot'],
        ['Appt No Show', '2:00pm, 2h delay', 'Reschedule after missed appointment'],
        ['Promotional', '11:00am, 24h delay', 'Limited offer, 50% off first month'],
    ],
    [1.5, 1.5, 3.0]
)

h3('AI Conversation Engine (Funnel Replies)')
body(
    'Every outbound funnel message has an AI sitting behind it. When a lead replies to any '
    'funnel SMS or email, the AI takes over the thread immediately and holds a full conversation. '
    'This replaces the silence that currently follows outbound messages and turns every reply '
    'into a live sales conversation running 24/7.'
)

body('How it works:')
for b in [
    'Lead replies to a funnel SMS or email',
    'Reply routes to Atlas via Twilio inbound webhook (SMS) or SMTP2GO inbound routing (email)',
    'Atlas identifies the lead, the campaign, the current deal stage, and the full conversation history',
    'Claude is invoked with full context: lead business info, stage, campaign objective, conversation history, and ZING knowledge base',
    'AI responds immediately and continues the conversation across as many exchanges as needed',
    'Every message in both directions is logged to the contact\'s CRM timeline',
]:
    bullet(b)

body('What the AI can do in a conversation:')
for b in [
    'Answer questions about ZING\'s products, pricing, and process',
    'Handle common objections (too expensive, already have a website, not the right time)',
    'Create urgency and move leads toward booking a call',
    'Book an appointment directly via Cal.com API when the lead agrees',
    'Update the deal stage automatically when a qualification signal is detected',
    'Unenroll the lead from the workflow if they opt out or are clearly disqualified',
    'Escalate to the assigned human rep when the lead specifically asks for a person, or when the AI detects a situation outside its confidence',
]:
    bullet(b)

body('Stage-specific AI objectives:')
table(
    ['Stage', 'AI Objective', 'Success Action'],
    [
        ['Call No Answer', 'Re-engage. Build rapport. Get commitment to a call.', 'Move deal to Active or Appointment'],
        ['Hot 72', 'Create urgency. Confirm interest. Book the appointment.', 'Move deal to Appointment, book via Cal.com'],
        ['Appt No Show', 'Acknowledge the miss without friction. Reschedule.', 'Move deal to Appointment, rebook via Cal.com'],
        ['Promotional', 'Answer offer questions. Qualify. Close on the spot if possible.', 'Move deal to Active or Won'],
    ],
    [1.5, 2.5, 2.0]
)

body('Human escalation rules:')
for b in [
    'Lead explicitly asks to speak to a person: AI replies "Let me get someone from our team to reach out shortly" and creates an internal task for the assigned rep',
    'AI confidence falls below threshold (ambiguous intent, complex objection, angry lead): flags for rep review without replying',
    'Lead is ready to buy but the AI cannot process payment: hands off to rep with full conversation context',
    'Rep takes over: AI disengages from the thread and all subsequent messages route to the rep directly',
]:
    bullet(b)

h3('Knowledge Base (AI Grounding)')
body(
    'The AI does not hallucinate answers. All responses are grounded in ZING\'s knowledge base, '
    'retrieved semantically at the time of each reply.'
)
for b in [
    'KB stored as vector embeddings in Supabase using the pgvector extension (no new service required)',
    'Content includes: product descriptions and pricing, onboarding process overview, common objections and responses, competitor comparisons, success stories, FAQs, and booking instructions',
    'At each reply, Atlas runs a semantic search against the KB to retrieve the most relevant context before invoking Claude',
    'KB is editable by admin users inside Atlas: add, update, or remove KB entries without touching code',
]:
    bullet(b)

# MODULE 10
doc.add_page_break()
h2('Module 10: Workflows')
body(
    'Visual automation builder. This is the engine that replaces HubSpot workflows entirely. '
    'Build multi-step automations triggered by any event in the platform, with branches, '
    'delays, AI calls, emails, SMS, and webhooks.'
)

h3('13 Trigger Types')
for b in [
    'Contact is created',
    'Deal stage changes',
    'Deal is won',
    'Payment fails (Stripe)',
    'Payment received (Stripe)',
    'Form is submitted',
    'Contact property changes',
    'Manual enrollment',
    'Date-based (scheduled)',
    'Onboarding task completed',
    'Onboarding task overdue',
    'Website stage changes',
    'All onboarding tasks complete',
]:
    bullet(b)

h3('9 Node Types')
table(
    ['Node', 'Description'],
    [
        ['Send Email', 'Compose and send an email to the contact with merge fields'],
        ['Send SMS', 'Send an SMS via Twilio with merge fields'],
        ['MAX Call (AI)', 'Trigger a Vapi AI outbound call with custom instructions and objective'],
        ['MAX Text (AI)', 'Deploy the full AI conversation engine on a thread. Handles ongoing back-and-forth, not just a single outbound message. Can book appointments, move deal stages, and escalate to humans.'],
        ['Send Webhook', 'POST to an external URL (e.g. collections agent, Zapier)'],
        ['Delay', 'Wait X days and Y hours before the next node executes'],
        ['If/Then Branch', 'Split the workflow based on a contact or deal property condition'],
        ['Edit Record', 'Update a contact or deal property automatically'],
        ['Manual Call', 'Assign a task for a team member to make a manual call'],
    ],
    [1.8, 4.2]
)

h3('8 Pre-Built Workflows (Ready at Launch)')
table(
    ['Workflow', 'Trigger', 'Purpose'],
    [
        ['AR Collections: Past Due', 'Payment fails (Stripe)', 'SMS > 3-day wait > email > 4-day wait > branch > MAX Call if still unpaid'],
        ['New Customer Onboarding', 'Deal is won', 'Welcome email > day 1 SMS > week 1 progress email > MAX check-in text'],
        ['Domain Renewal Check', 'Date-based', 'Branch on renewal status > 14-day wait > webhook > clear property'],
        ['Win-Back Campaign', 'Property change (cancelled)', '7-day wait > "We miss you" email > 14-day wait > discount offer > MAX Call'],
        ['Website First Draft Ready', 'Website stage: first-draft', 'Email with preview link > SMS notification > 3-day wait > MAX follow-up text'],
        ['Website Sent for Review', 'Website stage: website-sent', 'Review email > 5-day wait > MAX Call to check approval'],
        ['Onboarding Task Overdue', 'Onboarding task overdue', 'Internal team SMS alert > 2-day wait > escalation email > 3-day wait > branch > escalation to team lead'],
        ['Customer Fully Onboarded', 'All onboarding tasks complete', 'Congrats email + SMS > 7-day wait > MAX post-launch call > 14-day wait > 3-week results email'],
    ],
    [1.8, 1.2, 3.0]
)

# MODULE 11
h2('Module 11: Reporting')
body(
    'Full analytics across every department. Date range filtering with presets on every report. '
    'Built for Amy and managers to assess performance without exporting to spreadsheets.'
)

h3('Revenue Report')
for b in [
    'Total Revenue, Subscription Revenue, and One-Time Revenue for the period',
    'Average Transaction Value (ATV) and Average Spend per customer',
    'Active MRR, Estimated LTV, Customer Acquisition Cost (CAC)',
    'Churned Revenue, Churn Rate %, NRR %',
    'Unique customers acquired in period',
    'Revenue breakdown by product (DISCOVER, BOOST, DOMINATE)',
]:
    bullet(b)

h3('Employee Performance Report')
for b in [
    'Per-rep breakdown: Revenue, Units Sold, ATV, Cancelled Revenue, Cancelled Units',
    'Sortable leaderboard with drill-down into individual rep\'s deals',
    'Cancellation attribution: which rep\'s customers have churned and how much revenue was lost',
]:
    bullet(b)

h3('Onboarding Report')
for b in [
    'Per onboarding item: Total, Complete, In Progress, Not Started, Overdue, Completion %',
    'Owner breakdown per item: which team members are active vs complete on each task',
    'Completion rate trends over time',
]:
    bullet(b)

# MODULE 12
h2('Module 12: AI Insights')
body(
    'A dedicated intelligence layer that runs continuously across all Atlas data and surfaces '
    'actionable recommendations. Not a dashboard decoration. A genuine business advisor.'
)

h3('How It Works')
for b in [
    'Analysis jobs run nightly across all platform data',
    'Event-triggered jobs fire on key events: deal won, customer churned, payment failed, onboarding completed',
    'Claude (Anthropic AI) analyses patterns and generates plain-English recommendations',
    'Insights stored in the database and surfaced throughout the platform',
]:
    bullet(b)

h3('What It Analyses')
table(
    ['Area', 'Example Insight'],
    [
        ['Sales Performance', '"Eric\'s close rate dropped 23%. His hot-72 conversion is half the team average."'],
        ['Onboarding Health', '"Customers completing GBP Optimisation in week 1 have 40% lower 90-day churn."'],
        ['Financial', '"Revenue is $11,400 below MRR target. At current velocity, the gap closes by $7,200."'],
        ['Customer Health', '"3 dental practices churned in 60 days. Vertical-specific spike, investigate before acquiring more."'],
        ['Team Capacity', '"Raj has 23 active onboardings, which is above the safe capacity threshold."'],
        ['Marketing', '"Email campaigns produce 34% longer-retained customers than paid ads."'],
        ['Support Patterns', '"Login issues reported 5 times this month. This is systemic, not isolated."'],
    ],
    [1.8, 4.2]
)

h3('Where Insights Surface')
for b in [
    'Dedicated Insights page: full history with status (acted on, dismissed, or pending)',
    'Dashboard: top 3 insights today',
    'Contextual nudges inline throughout every module',
    'Weekly AI summary email to Amy (configurable)',
]:
    bullet(b)

# MODULE 13
h2('Module 13: Data Infrastructure')
body(
    'The most strategically significant component of Atlas. The aggregate, anonymised data '
    'across 100,000 small businesses has genuine commercial value. This must be architected '
    'correctly from day one.'
)

h3('Onboarding Data Capture')
body('Beyond what is needed to deliver service, every onboarding captures rich business intelligence:')
for b in [
    'Industry and business type (standardised taxonomy)',
    'Years in business, number of employees, annual revenue range',
    'Monthly marketing budget and current marketing channels in use',
    'Current tools and software: CRM, booking, POS, etc.',
    'Existing website, Google Business Profile, and social presence status',
    'How they heard about ZING and what problem they are solving',
    'Owner/operator age range and business ownership history',
    'Geographic market: city, state, metro area, market size classification',
]:
    bullet(b)

h3('Event Tracking')
for b in [
    'PostHog embedded on Atlas tracks every employee interaction, feature usage, and operational pattern',
    'Lightweight analytics pixel embedded on every customer website we build (injected by the Press agent on every Duda site)',
    'Pixel captures: page views, unique visitors, form submissions, conversions, traffic sources, bounce rates, device data, and geographic data',
    'Builds a dataset on what actually works for small business websites across every vertical ZING serves',
]:
    bullet(b)

h3('Churn Intelligence')
body('Every churn event is captured in full:')
for b in [
    'Exact date and stage of cancellation',
    'Onboarding completion status at time of churn',
    'Days since last login, last payment, last support contact',
    'Preceding events: payment failures, support tickets, onboarding delays',
    'Industry, geography, the rep who sold them, the designer who onboarded them',
    'Over time this builds a churn prediction model that flags at-risk accounts before they cancel',
]:
    bullet(b)

h3('The Data Asset: Monetisation Paths')
table(
    ['Data Product', 'Likely Buyer', 'Value'],
    [
        ['SMB financial health by vertical and geography', 'Banks, lenders, insurance', 'Credit risk modelling'],
        ['Website conversion benchmarks by industry', 'Marketing agencies, ad platforms', 'Campaign optimisation'],
        ['Churn pattern signals for SMBs', 'SaaS companies, investors', 'Retention tools'],
        ['Marketing channel ROI by business type', 'Ad platforms, agencies', 'Benchmarking'],
        ['Local business health index by metro area', 'Commercial real estate, PE firms', 'Market analysis'],
    ],
    [2.0, 1.8, 1.8]
)

note(
    'LEGAL REQUIREMENT: Before any data is sold or shared externally, Terms of Service and Privacy Policy '
    'must explicitly disclose data collection and third-party use of anonymised, aggregated data. '
    'Customers must provide informed consent during onboarding. Recommend legal review before launch. '
    'Data sold must be anonymised and aggregated, never traceable to an individual business.'
)

h3('Data Warehouse')
for b in [
    'Separate from the operational database, built purely for analytics',
    'Nightly ETL pipeline syncs operational data and event data into BigQuery',
    'Where trend analysis, cohort modelling, and churn prediction run',
    'Where the sellable data product is structured',
]:
    bullet(b)

# MODULE 14
h2('Module 14: Settings')
body('Admin-only configuration panel.')
for b in [
    'Team Members: add, edit, and deactivate employees, assign roles',
    'Designers: manage offshore, US, and publisher lists',
    'Campaigns: create and manage campaign records',
    'Roles and Permissions: configure what each role can see and edit',
    'Data Consent: manage consent records and data collection configuration',
    'Integrations: Stripe, Twilio, SMTP2GO, HubSpot connection status',
]:
    bullet(b)

# AUTH
h2('Auth and Access Control')
body(
    'Google OAuth restricted to @zing-work.com. Employees sign in with their Google account, '
    'no passwords to manage. Role-based access is enforced at the database level via '
    'Row-Level Security. A sales rep cannot query another rep\'s data. This is enforced '
    'by the database itself, not just the UI.'
)

table(
    ['Role', 'Access'],
    [
        ['Admin', 'Everything: full read and write across all modules'],
        ['Manager', 'Full read, own department writes'],
        ['Sales Rep', 'Own pipeline only, team leaderboard read-only'],
        ['Designer', 'Assigned onboarding queue only'],
        ['Publisher', 'Publishing queue only'],
        ['Finance', 'AR and billing module'],
        ['Support', 'Ticket queue only'],
        ['Analyst', 'Data and Insights read-only'],
    ],
    [1.5, 4.5]
)

# BUSINESS LOGIC
h2('Key Business Logic')
for b in [
    'Won deal > auto-create onboarding record with all 12 items and due dates',
    'Won deal > enroll contact in New Customer Onboarding workflow automatically',
    'Payment failure > trigger AR Collections workflow automatically',
    'NRR = (starting MRR + expansion MRR) divided by starting MRR, expressed as %',
    'MRR per business day excludes weekends',
    'Launch fee split payment validation: allocated amounts must equal total fee',
    'Commission calculation per product commission structure',
    'Onboarding due dates auto-calculated per item type from sale date',
    'Contact communication timeline aggregates pre-sale and post-sale data from all modules',
    'Churn risk scoring based on: payment behaviour, onboarding completion, engagement, and industry signals',
    'Pipeline funnel SMS fires when a deal enters the configured stage, after the configured delay',
    'AI analysis runs nightly and on event triggers, results stored and surfaced throughout Atlas',
    'Event data pipeline to BigQuery runs nightly',
]:
    bullet(b)

# INTEGRATIONS
h2('Integration Map')
table(
    ['Integration', 'Purpose', 'Phase'],
    [
        ['Supabase Auth + Google OAuth', 'Employee login', '1'],
        ['PostHog (Atlas platform)', 'Employee and operational event tracking', '1'],
        ['PostHog (customer sites)', 'Website analytics pixel via Press agent', '2'],
        ['Stripe webhooks', 'Auto-update AR status on payment events', '3'],
        ['Claude API (Anthropic)', 'AI analysis and recommendations engine', '3'],
        ['Twilio (outbound + inbound)', 'AR SMS, workflow SMS, and inbound reply routing to AI engine', '3'],
        ['SMTP2GO (outbound + inbound)', 'Email notifications, workflow emails, weekly AI summary, and inbound reply routing to AI engine', '3'],
        ['Cal.com API', 'AI books appointments directly inside lead conversations', '3'],
        ['Supabase pgvector', 'Knowledge base for AI conversation grounding (no new service)', '3'],
        ['BigQuery', 'Data warehouse for analytics and monetisation', '4'],
        ['HubSpot (read-only)', 'One-time data migration', '4'],
        ['Scout / Catch agents', 'Re-pointed to write to Atlas instead of HubSpot', '4'],
        ['Press agent', 'Updates onboarding record when site is published', '4'],
    ],
    [2.5, 2.5, 0.6]
)

# PHASE PLAN
h2('6-Week Build Plan')
table(
    ['Week', 'Focus', 'Deliverable'],
    [
        ['1', 'Foundation', 'Supabase schema, auth, Next.js scaffold, API on Railway, seed data, roles'],
        ['2', 'Core Modules', 'Dashboard, Pipeline, Contacts, Products wired to live API'],
        ['3', 'Full Platform', 'Onboarding (both views), AR, Tickets, Campaigns, Workflows, Reporting, Settings'],
        ['4', 'AI + Data Layer', 'Claude analysis jobs, PostHog events, churn risk scores, Insights page, pixel on Duda sites'],
        ['5', 'Integrations', 'Stripe webhooks, Twilio, SMTP2GO, Scout/Catch/Press re-pointing'],
        ['6', 'Migration + Polish', 'HubSpot data import, UAT with Amy, performance tuning, go-live'],
    ],
    [0.6, 1.5, 3.5]
)

# ============================================================
# PART 2: TECH STACK
# ============================================================
doc.add_page_break()
h1('Part 2: Tech Stack')

body(
    'Every decision in this stack is driven by five principles: cloud-first (not dependent on '
    'one machine staying on), standard tools that every engineer knows, one platform where '
    'possible to reduce vendors, built for data from day one, and AI-ready throughout.'
)

h2('Frontend: Next.js on Vercel')
for b in [
    'Industry standard for internal dashboards at this scale',
    'Proper page routing with each module having its own URL',
    'Server-side rendering for fast initial loads even with large datasets',
    'Code splitting keeps the app fast as it grows',
    'TypeScript for reliability and maintainability',
    'Tailwind CSS for consistent, maintainable design system',
    'Recharts or Tremor for charts and dashboard components',
]:
    bullet(b)

body(
    'Hosted on Vercel. Zero-configuration automatic deployments. Push to GitHub and it deploys. '
    'Global CDN. Free tier. If the Mac Studio goes offline, employees can still open Atlas.'
)

note('The single-file CDN React approach from the prototype is a prototyping pattern, not a production pattern for 100 employees. Next.js is the correct choice at this scale.')

h2('Database: Supabase (PostgreSQL)')
body('Supabase combines four services into one platform, eliminating multiple vendors:')

table(
    ['Need', 'What Supabase Provides'],
    [
        ['PostgreSQL database (100,000+ customers, millions of records)', 'Built in'],
        ['Google OAuth authentication', 'Built in'],
        ['Row-Level Security (role-based data access enforced at DB level)', 'Built in'],
        ['Realtime subscriptions (live dashboard updates across all users)', 'Built in'],
        ['File and document storage', 'Built in'],
        ['Admin data dashboard', 'Built in'],
        ['Auto-generated REST API for simple CRUD operations', 'Built in'],
    ],
    [3.5, 2.5]
)

body(
    'PostgreSQL handles 100,000 customers without strain. Real-world ceiling is tens of millions '
    'of records. Row-Level Security means a sales rep literally cannot query another rep\'s deals, '
    'enforced by the database itself, not just the application layer.'
)
body('Cost: Free tier for development. $25/month Pro for production.')

h2('API Layer: Node.js + Fastify on Railway')
body(
    'Handles complex business logic beyond what Supabase\'s auto-generated API covers: '
    'won deal triggering onboarding creation, HubSpot data migration, Stripe webhook processing, '
    'AI job triggers, and complex calculations.'
)
for b in [
    'Fastify over Express: faster, better TypeScript support, same ecosystem and patterns',
    'Railway: simple cloud hosting, auto-deploy on GitHub push, $5 to $10 per month',
    'No Mac Studio dependency. The API stays up independently.',
    'Existing ZING agents (Scout, Press, Chase) point to this API with no pattern change required',
]:
    bullet(b)

h2('Event Tracking: PostHog')
body(
    'Open source, privacy-first event tracking. Two deployments: one on Atlas for employee '
    'interaction data, one as a pixel on every customer website we build.'
)
for b in [
    'Tracks every employee interaction, feature usage, and operational pattern on Atlas',
    'Lightweight analytics pixel injected by the Press agent on every Duda customer site we build',
    'Pixel captures: page views, unique visitors, form submissions, conversions, traffic sources, bounce rates, device and geographic data',
    'Includes session recording, funnels, cohort analysis, and feature flags',
    'Free tier: 1 million events per month, more than sufficient for initial scale',
    'Feeds directly into BigQuery as part of the nightly data pipeline',
]:
    bullet(b)

h2('Data Warehouse: BigQuery (Google Cloud)')
body(
    'Separate from the operational database. Built purely for analytics, trend analysis, '
    'and the data monetisation layer.'
)
for b in [
    'Google\'s serverless analytics warehouse with no infrastructure to manage',
    'Free tier: 10GB storage and 1TB queries per month, sufficient for years of growth',
    'Scales to petabytes with no architectural changes required',
    'Purpose-built for the cohort analysis, churn modelling, and trend detection the data asset requires',
    'Native connectors to PostHog, Supabase, and Looker Studio for visualisation',
    'BigQuery ML for churn prediction models as the dataset matures',
]:
    bullet(b)

body('A nightly ETL job syncs operational data from Supabase and event data from PostHog into BigQuery.')

h2('AI Intelligence: Claude API + Trigger.dev')

h3('Claude API (Anthropic)')
for b in [
    'ZING already has an active API key',
    'Primary model for all analysis, recommendations, and natural language insights',
    'Claude Sonnet for routine nightly analysis (cost-efficient)',
    'Claude Opus for deep weekly reports and complex pattern detection',
]:
    bullet(b)

h3('Trigger.dev: Background Job Platform')
body('Designed specifically for AI workflows. Handles:')
for b in [
    'Nightly analysis jobs (run at 2am, analyse all data, generate and store insights)',
    'Event-triggered jobs (deal marked won > trigger onboarding risk analysis > update churn scores)',
    'Long-running jobs without timeout issues',
    'Free tier handles this volume comfortably. $50/month Pro tier at scale.',
    'Runs in the cloud, independent of Mac Studio uptime',
]:
    bullet(b)

h2('AI Conversation Engine: Claude + pgvector + Twilio Inbound + SMTP2GO Inbound')
body(
    'The AI conversation layer sits behind every outbound marketing message. '
    'It requires three additions to the base stack, none of which introduce new vendors.'
)

h3('Inbound SMS: Twilio Inbound Webhook')
for b in [
    'Twilio is already in the stack for outbound SMS',
    'Adding inbound is a single config change: point the Twilio Messaging Service inbound webhook to the Atlas API',
    'Atlas receives the reply, identifies the lead and thread, invokes Claude with full context, sends the response back via Twilio',
    'No new service required',
]:
    bullet(b)

h3('Inbound Email: SMTP2GO Inbound Routing')
for b in [
    'SMTP2GO is already in the stack for outbound email',
    'SMTP2GO supports inbound routing: replies to a configured address (e.g. ai@zing-work.com) are POSTed to an Atlas webhook',
    'Atlas receives the reply, identifies the lead and thread, invokes Claude, sends the response via SMTP2GO',
    'No new service required',
]:
    bullet(b)

h3('Knowledge Base: Supabase pgvector')
for b in [
    'pgvector is a Supabase extension, already available on all plans at no extra cost',
    'Stores ZING\'s KB documents as vector embeddings (product info, pricing, objections, FAQs, process)',
    'At each AI reply, Atlas performs a semantic similarity search to retrieve the most relevant KB chunks',
    'Retrieved context is injected into the Claude prompt to ground the response and prevent hallucination',
    'KB management UI inside Atlas Settings allows admin users to add, edit, and delete entries',
]:
    bullet(b)

h3('Appointment Booking: Cal.com API')
for b in [
    'Cal.com is already used for Amy\'s partnership booking link',
    'Adding API access allows the AI to check rep availability and create bookings directly inside a conversation',
    'When a lead says "yes let\'s set up a call", the AI books it without a human in the loop',
    'Booking confirmation sent to the lead automatically via SMS and email',
    'Cal.com API access: free on the existing plan',
]:
    bullet(b)

h3('Conversation State: New Supabase Tables')
for b in [
    'conversation_threads: one row per lead+channel thread (SMS or email), stores status and assigned rep',
    'conversation_messages: every inbound and outbound message stored with timestamp, direction, and content',
    'No new infrastructure. Two new tables in the existing Supabase database.',
]:
    bullet(b)

h2('Realtime: Supabase Realtime')
body(
    'Built into Supabase at no extra cost. When a deal is closed, the dashboard updates for '
    'all users instantly. When onboarding status changes, the board updates live across all '
    'designers simultaneously.'
)

h2('Auth: Supabase Auth')
for b in [
    'Google OAuth built in. Employees sign in with their @zing-work.com Google account.',
    'Restricted to the zing-work.com domain only',
    'No passwords to manage, no reset flows to build',
    'JWT session management handled automatically',
    'Role stored in the database and enforced by Row-Level Security',
]:
    bullet(b)

h2('CI/CD: GitHub')
body(
    'All code lives on GitHub. Vercel auto-deploys the frontend on every push to main. '
    'Railway auto-deploys the API. Push code and it is live in 60 seconds. '
    'No manual deployments, no SSH, no clasp push.'
)

h2('Full Stack at a Glance')
table(
    ['Layer', 'Technology', 'Hosted On', 'Est. Monthly Cost'],
    [
        ['Frontend', 'Next.js (React + TypeScript)', 'Vercel', 'Free to $20'],
        ['Auth', 'Supabase Auth + Google OAuth', 'Supabase', 'Included'],
        ['Operational Database', 'PostgreSQL', 'Supabase', 'Free to $25'],
        ['Realtime', 'Supabase Realtime', 'Supabase', 'Included'],
        ['API / Business Logic', 'Node.js + Fastify', 'Railway', '$5 to $10'],
        ['Event Tracking', 'PostHog', 'PostHog Cloud', 'Free to usage'],
        ['Data Warehouse', 'BigQuery', 'Google Cloud', 'Free to $5'],
        ['AI Engine', 'Claude API (Anthropic)', 'Anthropic', '$20 to $50'],
        ['Background Jobs', 'Trigger.dev', 'Trigger.dev Cloud', 'Free to $50'],
        ['AI Conversation Engine', 'Claude API + pgvector + Twilio Inbound + SMTP2GO Inbound', 'Existing services', 'Included in existing costs'],
        ['Knowledge Base', 'Supabase pgvector extension', 'Supabase', 'Included'],
        ['Appointment Booking (AI)', 'Cal.com API', 'Cal.com', 'Included in existing plan'],
        ['CI/CD', 'GitHub Actions', 'GitHub', 'Free'],
    ],
    [1.8, 1.8, 1.5, 1.1]
)

h2('Why This Stack vs Alternatives')
table(
    ['Alternative Considered', 'Why Not Chosen'],
    [
        ['Single-file CDN React', 'Prototype pattern only. No routing, no code splitting, not suitable for 100 employees.'],
        ['SQLite', 'File-level locking. Cannot support concurrent writes from 100 users.'],
        ['Mac Studio as only host', 'Single point of failure. One machine offline means 100 employees are down.'],
        ['Firebase / Firestore', 'NoSQL document store. Poor fit for relational CRM data and financial reporting.'],
        ['Custom auth from scratch', 'Security risk and wasted build time. Supabase Auth solves this completely.'],
        ['No event tracking', 'The data asset cannot be built without capturing events from day one.'],
        ['Polling instead of realtime', 'Poor UX for 100 simultaneous users. Supabase Realtime is included at no extra cost.'],
    ],
    [2.3, 3.7]
)

# SETUP
doc.add_page_break()
h2('What Is Needed to Start Building')
body('Five accounts to set up before Atlas begins. All have free tiers.')
table(
    ['Service', 'URL', 'What to Create'],
    [
        ['Supabase', 'supabase.com', 'New project named "atlas". Send the project URL and anon key.'],
        ['Railway', 'railway.app', 'Create account and connect the ZING GitHub organisation.'],
        ['PostHog', 'posthog.com', 'Create account. Send the project API key.'],
        ['Google Cloud', 'console.cloud.google.com', 'Enable BigQuery. Create dataset named "zing_analytics".'],
        ['GitHub', 'github.com/zinglocal', 'Create repo named "atlas" under the ZING organisation.'],
    ],
    [1.2, 1.8, 3.0]
)

note(
    'Legal: Before the data collection and monetisation features go live, Terms of Service and '
    'Privacy Policy must be reviewed and updated by a lawyer to explicitly disclose data collection '
    'and third-party use of anonymised, aggregated data. Estimated cost: $500 to $2,000. '
    'This protects the data asset, which is the most valuable long-term component of Atlas.'
)

divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Atlas  |  The ZING Operating Platform  |  Confidential  |  Prepared by Max  |  March 2026')
run.font.size = Pt(8)
run.font.color.rgb = GREY

out = '/Users/zingadmin/Downloads/Atlas_ZING_Operating_Platform.docx'
doc.save(out)
print(f'Saved: {out}')
