from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

OXFORD = RGBColor(0x05, 0x05, 0x36)
PURPLE = RGBColor(0x64, 0x07, 0xFA)
GREY   = RGBColor(0x5A, 0x5F, 0x7A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)
ORANGE = RGBColor(0xC0, 0x55, 0x10)
GREEN  = RGBColor(0x05, 0x96, 0x68)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = OXFORD
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '6407FA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PURPLE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = OXFORD
    return p

def h4(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREY
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p

def note(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run('  ' + text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = color or ORANGE
    run.italic = True
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '050536')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = 'F5F7FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(chr(0x2500) * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def section_tag(text, color_hex):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'  {text}  ')
    run.bold = True
    run.font.size = Pt(9)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)

# ============================================================
# COVER
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('ATLAS')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FEATURE SPECIFICATION')
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = OXFORD

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MAX AI  |  Data Hub')
run.font.size = Pt(13)
run.font.color.rgb = GREY
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Prepared by Max, Chief of Staff  |  March 2026  |  CONFIDENTIAL')
run.font.size = Pt(9)
run.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# MODULE: MAX AI
# ============================================================
h1('Module: MAX AI')

body(
    'MAX AI is the intelligence and conversation center for Atlas. It gives Amy and her team '
    'full visibility into what the AI is doing across the business: the insights it has generated, '
    'the conversations it is having with leads, the knowledge base it draws on, and the performance '
    'metrics that show whether it is working. MAX AI is also where the AI is configured and tuned.'
)

table(
    ['Attribute', 'Value'],
    [
        ['Left nav label', 'MAX AI'],
        ['Left nav position', 'Bottom section of nav (with Reporting and Data Hub)'],
        ['Icon', 'Spark / lightning bolt'],
        ['Access', 'Admin and Manager: all tabs. Sales Rep: Conversations tab (own leads only).'],
        ['Tabs', '5: Insights, Conversations, Knowledge Base, Performance, Configuration'],
    ],
    [2.0, 4.0]
)

# TAB 1: INSIGHTS
h2('Tab 1: Insights')
body(
    'All AI-generated business recommendations in one place. Every insight has a status so the '
    'team can track whether action was taken. This is the primary surface through which the '
    'nightly analysis engine communicates with Amy and her managers.'
)

h3('Stat Cards (top row)')
for b in [
    'New Insights This Week',
    'Insights Acted On (lifetime count)',
    'Insights Dismissed (lifetime count)',
    'Avg Time to Action (how quickly the team acts on AI recommendations)',
]:
    bullet(b)

h3('Insight Cards')
body('Each insight is displayed as a card in a scrollable list. Cards contain:')
for b in [
    'Category badge: Sales, Onboarding, Financial, Customer, Marketing, or Support',
    'Severity badge: Critical (red), Important (orange), or Informational (grey)',
    'Headline finding in plain English (e.g. "Raj\'s onboarding capacity is above safe threshold")',
    'Detail paragraph with the data points that support the finding',
    'Recommended action',
    'Date generated',
    'Status: New, Acted On, or Dismissed',
    'Actions: Mark Acted On, Dismiss, Share (sends to a specific team member via in-app notification)',
]:
    bullet(b)

h3('Filters')
for b in [
    'Category (multi-select)',
    'Severity',
    'Status (New / Acted On / Dismissed)',
    'Date range',
]:
    bullet(b)

h3('Insight Detail Panel (side panel on click)')
for b in [
    'Full analysis text with supporting data points',
    'Trend chart where relevant (e.g. rep performance over 90 days)',
    'Notes field for the team to log what action was taken',
    'Linked records: click through to the relevant deal, contact, or onboarding record',
]:
    bullet(b)

# TAB 2: CONVERSATIONS
h2('Tab 2: Conversations')
body(
    'All AI-driven lead conversations across SMS and email. This is where the team monitors '
    'active threads, reviews escalations, and takes over conversations when needed.'
)

h3('Stat Cards (top row)')
for b in [
    'Active Conversations Now',
    'Conversations This Month',
    'Appointment Booking Rate (% of conversations that resulted in a booking)',
    'Stage Conversion Rate (% of conversations that moved the deal forward)',
    'Escalation Rate (% of conversations handed to a human)',
    'Avg Response Time (seconds from inbound message to AI reply)',
]:
    bullet(b)

h3('Escalation Queue (prominent, shown only when pending)')
body(
    'A highlighted panel at the top of the tab when any conversations are flagged for human review. '
    'Shown above the main table. Each escalation shows the lead name, reason for escalation, '
    'and a Claim button that assigns the thread to the logged-in rep.'
)
body('Escalation reasons:')
for b in [
    'Lead asked to speak to a person',
    'AI confidence fell below the configured threshold',
    'Lead is ready to buy but AI cannot process payment',
    'Conversation has been active for more than the configured timeout with no resolution',
]:
    bullet(b)

h3('Conversations Table')
body('Columns: Lead Name, Business, Channel (SMS / Email), Deal Stage, Message Count, Last Activity, Status')

body('Status badge colours:')
for b in [
    'Active: green',
    'Resolved: grey',
    'Escalated: orange (waiting for human review)',
    'Human Takeover: blue (a rep has taken over the thread)',
]:
    bullet(b)

h3('Conversation Detail (click into a thread)')
for b in [
    'Full chronological message thread showing every inbound and outbound message',
    'Each message shows: direction (Inbound / Outbound AI / Outbound Human), timestamp, channel, content',
    'Outbound AI messages include a collapsible panel showing which KB chunks were retrieved to ground that response (for admin quality review)',
    'Take Over button: assigns thread to logged-in rep, AI disengages immediately, all future replies route to rep',
    'Re-engage AI button (admin only): returns thread to AI control',
    'Move Deal Stage: manually advance or reverse the deal stage from inside the conversation',
    'Add Note: internal note visible to the team only, not sent to the lead',
]:
    bullet(b)

# TAB 3: KNOWLEDGE BASE
h2('Tab 3: Knowledge Base')
body(
    'The content the AI draws on to ground every response. Editable by admin users '
    'without touching code. Every entry is stored as a vector embedding so the AI '
    'can retrieve the most relevant content for any question a lead asks.'
)

h3('Stats')
for b in [
    'Total KB entries',
    'Entries indexed (vectors ready)',
    'Entries pending indexing',
    'Last indexed timestamp',
]:
    bullet(b)

h3('KB Categories')
table(
    ['Category', 'What It Contains'],
    [
        ['Products', 'DISCOVER, BOOST, DOMINATE descriptions. What is included. What is not included.'],
        ['Pricing', 'Current prices, any active discounts or promotions, upgrade paths'],
        ['Objections', 'Common lead pushbacks and how to handle them (too expensive, already have a website, not the right time, bad experience with agencies)'],
        ['Process', 'How onboarding works, timelines, what the customer can expect at each stage'],
        ['Competitors', 'How ZING compares to Wix, GoDaddy, Squarespace, Duda DIY, and local agencies'],
        ['FAQs', 'Anything a lead might ask that does not fit another category'],
        ['Booking', 'How to schedule a call, what happens on the call, who they will speak to, Cal.com link'],
    ],
    [1.5, 4.5]
)

h3('KB Entry Table')
body('Columns: Title, Category, Last Updated, Status (Indexed / Pending / Error)')
body('Actions per row: Edit, Delete, Re-index')

h3('Add / Edit Entry Modal')
for b in [
    'Fields: Title, Category (dropdown), Content (rich text), Tags (optional)',
    'On save, entry is queued for vector indexing',
    'Status updates to Indexed automatically once complete (typically under 30 seconds)',
]:
    bullet(b)

h3('Test KB Search')
body(
    'A search input at the top of the tab. Type any question a lead might ask. '
    'Returns the top 3 KB chunks that would be retrieved for that query, ranked by relevance score. '
    'Allows admin to verify the KB is returning the right content before a real conversation depends on it.'
)

# TAB 4: PERFORMANCE
h2('Tab 4: Performance')
body('AI performance analytics. Date range filtered with presets. Broken down by channel and stage.')

h3('Conversation Metrics')
for b in [
    'Total conversations by channel (SMS vs Email)',
    'Avg message count per conversation',
    'Appointment booking rate overall and by stage',
    'Stage conversion rate by stage (which stages does the AI convert best from)',
    'Escalation rate by stage (where does the AI struggle most)',
    'Opt-out rate (leads who disengaged or unsubscribed during an AI conversation)',
]:
    bullet(b)

h3('Charts')
for b in [
    'Conversations over time: line chart, last 90 days, split by SMS and email',
    'Conversion funnel: Conversations Started > Replied > Qualified > Booked > Won',
    'Escalation reasons breakdown: pie chart showing proportion of each escalation type',
]:
    bullet(b)

h3('KB Performance')
for b in [
    'Most frequently retrieved KB chunks (what topics come up most in real conversations)',
    'Low-relevance flags: chunks that are being retrieved but scoring poorly, flagged for admin review and improvement',
]:
    bullet(b)

# TAB 5: CONFIGURATION
h2('Tab 5: Configuration')
body('Admin only. All settings for the AI engine.')

h3('Model Settings')
for b in [
    'Conversation model: dropdown (Claude Sonnet / Claude Opus)',
    'Nightly analysis model: dropdown (Claude Sonnet / Claude Opus)',
    'Max tokens per conversation reply',
    'Retrieval chunk count: how many KB chunks to inject per reply (default 3)',
]:
    bullet(b)

h3('Escalation Rules')
for b in [
    'Confidence threshold: if AI confidence falls below X%, escalate instead of replying (default 70%)',
    'Escalation keywords: list of phrases that always trigger human handoff (e.g. "speak to a person", "call me", "I want to talk to someone")',
    'Auto-escalate after X hours of active conversation with no resolution (default 48 hours)',
]:
    bullet(b)

h3('Working Hours')
for b in [
    'AI active hours (default 8:00am to 10:00pm)',
    'Timezone configuration',
    'Outside-hours behaviour: queue replies until active, or respond immediately with a "we will be in touch" message',
]:
    bullet(b)

h3('Nightly Analysis Schedule')
for b in [
    'Analysis run time (default 2:00am)',
    'Toggle individual analysis categories on or off: Sales, Onboarding, Financial, Customer, Marketing, Support',
    'Weekly AI summary email: toggle on/off, recipient email address',
]:
    bullet(b)

h3('Notification Settings')
for b in [
    'Who receives a notification when a new escalation is flagged (by role or specific team member)',
    'Escalation notification channel: in-app, SMS, or email',
]:
    bullet(b)

h2('MAX AI: Data Model')
body('New database tables required:')
table(
    ['Table', 'Purpose', 'Key Columns'],
    [
        ['ai_insights', 'Stores all AI-generated insights', 'id, category, severity, headline, detail, recommendation, status, generated_at, acted_at, acted_by'],
        ['conversation_threads', 'One row per lead+channel thread', 'id, contact_id, channel, deal_id, status, assigned_rep_id, created_at, last_activity_at'],
        ['conversation_messages', 'Every message in every thread', 'id, thread_id, direction, content, timestamp, kb_chunks_used (jsonb), ai_confidence (float)'],
        ['kb_entries', 'Knowledge base content with vector embeddings', 'id, title, category, content, tags, embedding (vector), indexed_at, status'],
        ['ai_config', 'AI configuration settings (single row)', 'model settings, escalation rules, working hours, schedule config'],
    ],
    [1.6, 1.8, 2.6]
)

# ============================================================
# MODULE: DATA HUB
# ============================================================
doc.add_page_break()
h1('Module: Data Hub')

body(
    'The Data Hub is the data asset management center for Atlas. It gives Amy and her team '
    'visibility into what data is being collected, from where, and at what volume. Over time, '
    'as the dataset grows across tens of thousands of customers, this module becomes the '
    'interface through which the data product is managed, monitored, and eventually surfaced '
    'to external buyers.'
)

note(
    'All aggregate data shown in the Data Hub is anonymised. No individual business is '
    'identifiable in any chart, benchmark, or export from the Aggregate Insights tab. '
    'The Consent Management tab is the source of truth for legal compliance.'
)

table(
    ['Attribute', 'Value'],
    [
        ['Left nav label', 'Data Hub'],
        ['Left nav position', 'Bottom section of nav (with Reporting and MAX AI)'],
        ['Icon', 'Database / stacked cylinders'],
        ['Access', 'Admin and Analyst roles only'],
        ['Tabs', '5: Overview, Business Profiles, Website Analytics, Aggregate Insights, Consent Management'],
    ],
    [2.0, 4.0]
)

# TAB 1: OVERVIEW
h2('Tab 1: Overview')
body('At-a-glance health of the data asset. The first thing an admin sees when opening the Data Hub.')

h3('Stat Cards (top row)')
for b in [
    'Total Customer Profiles (count of enriched customer records)',
    'Enrichment Completion Rate (% of customers with full onboarding data filled in)',
    'Events Collected Today (PostHog events across the Atlas platform and all customer sites combined)',
    'Last Warehouse Sync (timestamp and record count from the last BigQuery ETL run)',
    'Consented Customers (% and count of customers who have given data collection consent)',
]:
    bullet(b)

h3('Data Sources Panel')
body(
    'A status board showing each data source with its collection status, '
    'volume in the last 30 days, and last activity timestamp.'
)
table(
    ['Source', 'What It Captures'],
    [
        ['Onboarding Enrichment', 'Business profile fields captured during customer onboarding'],
        ['Atlas Platform Events', 'Every employee interaction and operational event on the Atlas platform (via PostHog)'],
        ['Customer Website Pixels', 'Traffic, conversions, and engagement data from every Duda site we have built (via PostHog pixel)'],
        ['Stripe Payment Events', 'Payment success, failure, and retry events linked to customer records'],
        ['Support Tickets', 'Ticket categories, resolution times, and repeat issue patterns'],
        ['Churn Events', 'Full churn records: timing, stage, preceding events, business context'],
    ],
    [2.2, 3.8]
)

h3('Warehouse Sync Log')
body('Last 10 BigQuery ETL sync runs showing: timestamp, duration, records written, and status (Success / Failed / Partial).')

# TAB 2: BUSINESS PROFILES
h2('Tab 2: Business Profiles')
body(
    'The enrichment data captured during onboarding for every customer. '
    'This is the core of the data asset. The richer and more complete this data, '
    'the more valuable the aggregate product becomes.'
)

h3('Table View')
body('Columns: Business Name, Industry, Years in Business, Employees, Revenue Range, Marketing Budget, State, Completion %')
body('Sortable, searchable, filterable by: industry, state, employee range, revenue range, completion status.')

h3('Field Completion Panel')
body(
    'A sidebar panel showing the completion rate for each enrichment field across all customers. '
    'Fields with low completion are flagged with a recommended fix.'
)
body('Example:')
for b in [
    'Industry: 94%',
    'Years in Business: 87%',
    'Revenue Range: 61%',
    'Marketing Budget: 48%',
    'Current Tools Used: 39% (flagged: consider making required at onboarding)',
]:
    bullet(b)

h3('Customer Profile Detail (click into a record)')
for b in [
    'Full business profile card with all enrichment fields displayed',
    'Consent status and consent date',
    'Link to the full CRM contact record',
    'Website analytics summary (if pixel is installed on their site)',
]:
    bullet(b)

h3('Enrichment Fields Captured')
table(
    ['Field', 'Captured At'],
    [
        ['Industry / business type (standardised taxonomy)', 'Onboarding'],
        ['Years in business', 'Onboarding'],
        ['Number of employees', 'Onboarding'],
        ['Annual revenue range', 'Onboarding'],
        ['Monthly marketing budget', 'Onboarding'],
        ['Current marketing channels in use', 'Onboarding'],
        ['Current tools and software (CRM, booking, POS, etc.)', 'Onboarding'],
        ['Existing website, GBP, and social presence status', 'Onboarding'],
        ['How they heard about ZING', 'Onboarding'],
        ['Primary business challenge they are solving', 'Onboarding'],
        ['Owner age range and business ownership history', 'Onboarding'],
        ['Geographic market: city, state, metro area, market size', 'Onboarding + CRM'],
    ],
    [3.5, 2.5]
)

# TAB 3: WEBSITE ANALYTICS
h2('Tab 3: Website Analytics')
body(
    'Per-customer website performance data collected via the PostHog analytics pixel '
    'embedded on every Duda site built by ZING. This data belongs to ZING, not the customer, '
    'and forms a key part of the aggregate data asset.'
)

h3('Table View')
body('Columns: Business Name, Site URL, Monthly Visitors, Avg Conversion Rate, Top Traffic Source, Bounce Rate, Last Updated')

h3('Customer Analytics Detail (click into a record)')
for b in [
    'Traffic by source: organic, paid, referral, direct, social (bar chart)',
    'Conversion funnel: visits > engaged sessions > form submissions > calls (funnel chart)',
    'Top pages by traffic volume',
    'Device split: desktop vs mobile vs tablet',
    'Geographic visitor breakdown by state',
    '30-day trend chart for visitors and conversions',
]:
    bullet(b)

h3('Aggregate Summary Panel')
for b in [
    'Average monthly visitors across all customer sites',
    'Average conversion rate by industry (top 10 industries, bar chart)',
    'Most common top traffic source across the portfolio',
    'Sites with pixel active vs pending installation (with action to notify the Press agent)',
]:
    bullet(b)

# TAB 4: AGGREGATE INSIGHTS
h2('Tab 4: Aggregate Insights')
body(
    'What the data asset looks like when anonymised and aggregated across all customers. '
    'This is the commercial data product in preview form. Over time, this tab becomes the '
    'most valuable page in the entire platform.'
)

note('All data on this tab is anonymised and aggregated. No individual business is identifiable in any chart or table.', GREEN)

h3('Industry Benchmarks')
for b in [
    'Average website conversion rate by industry (bar chart, top 15 industries)',
    'Average monthly visitors by industry',
    'Churn rate by industry: which verticals cancel most and least',
    'Average customer lifetime by industry (months)',
    'Average MRR at time of churn by industry',
]:
    bullet(b)

h3('Geographic Health')
for b in [
    'Choropleth map of the US showing customer density by state (colour intensity = customer count)',
    'Overlay toggle: switch between density view and health score view (health score = churn rate + payment behaviour)',
    'Click a state to see aggregate metrics for that market: customer count, avg MRR, churn rate, top industry',
]:
    bullet(b)

h3('Churn Intelligence')
for b in [
    'Average days from sale to cancellation by product (DISCOVER, BOOST, DOMINATE)',
    'Top 5 onboarding items incomplete at time of churn (bar chart)',
    'Churn rate trend over the last 12 months (line chart)',
    'Churn rate by industry (sortable table)',
    'Churn by lead source (which acquisition channels produce the shortest-lived customers)',
]:
    bullet(b)

h3('Marketing Effectiveness')
for b in [
    'Lead source to close rate: Email vs SMS vs Paid vs Referral vs Organic (horizontal bar chart)',
    'Average deal value by lead source',
    'Average time from first contact to close by lead source',
    'Customer lifetime by lead source (which channels produce the longest-retained customers)',
]:
    bullet(b)

# TAB 5: CONSENT MANAGEMENT
h2('Tab 5: Consent Management')
body(
    'The legal compliance layer for the data collection programme. '
    'Every consent record is stored here with full audit trail.'
)

h3('Table View')
body('Columns: Business Name, Contact, Consent Given (yes/no), Consent Date, Privacy Policy Version, Method (onboarding form / manual)')
body('Filterable by: consented, not consented, privacy policy version.')

h3('Actions')
for b in [
    'Export consent records as CSV (for legal compliance documentation and audit)',
    'Flag customers missing consent for outreach (generates a task list)',
    'Bulk mark as consented (for customers onboarded before consent tracking was implemented, with manual confirmation)',
]:
    bullet(b)

h3('Privacy Policy Version Tracker')
for b in [
    'When the privacy policy is updated, a new version is registered here',
    'Shows which customers consented to each version',
    'Flags customers who consented to an old version if re-consent is required under the new version',
    'Sends re-consent outreach via the Workflows engine when triggered',
]:
    bullet(b)

h2('Data Hub: Data Model')
body('New database tables required:')
table(
    ['Table', 'Purpose', 'Key Columns'],
    [
        ['data_consent', 'Customer consent records', 'id, customer_id, consent_given, consent_date, policy_version, method'],
        ['website_analytics_daily', 'Daily snapshot of each customer site\'s analytics', 'id, customer_id, date, visitors, conversions, bounce_rate, top_source, device_split (jsonb)'],
        ['platform_events', 'All Atlas platform events synced from PostHog', 'id, event_type, user_id, properties (jsonb), timestamp, source'],
        ['warehouse_sync_log', 'Log of each BigQuery ETL sync run', 'id, started_at, completed_at, records_written, status, error_detail'],
    ],
    [1.8, 2.0, 2.2]
)

body('Business enrichment fields are added as columns on the existing contacts table, not a separate table. This keeps the data co-located with the customer record and simplifies queries.')

# FOOTER
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Atlas  |  MAX AI + Data Hub Feature Specification  |  Confidential  |  Prepared by Max  |  March 2026')
run.font.size = Pt(8)
run.font.color.rgb = GREY

out = '/Users/zingadmin/Downloads/Atlas_MAX_AI_and_Data_Hub_Spec.docx'
doc.save(out)
print(f'Saved: {out}')
